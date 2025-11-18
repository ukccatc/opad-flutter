#!/usr/bin/env python3
"""
Generate readable titles for files by reading PDF/DOC content or improving filenames
"""
import os
import re
from pathlib import Path

def improve_filename(filename):
    """Improve filename to readable title"""
    # Remove extension
    name = filename.rsplit('.', 1)[0]
    
    # Replace common separators with spaces
    name = name.replace('-', ' ')
    name = name.replace('_', ' ')
    name = name.replace('№', '№ ')
    
    # Clean up multiple spaces
    name = re.sub(r'\s+', ' ', name)
    
    # Capitalize first letter of each word
    words = name.split()
    capitalized_words = []
    for word in words:
        if word.isupper() and len(word) > 1:
            # Keep acronyms as is
            capitalized_words.append(word)
        else:
            # Capitalize first letter
            capitalized_words.append(word.capitalize())
    
    return ' '.join(capitalized_words)

def extract_pdf_title(pdf_path):
    """Try to extract title from PDF metadata or first page"""
    try:
        # Try PyPDF2
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.metadata and reader.metadata.get('/Title'):
                    title = reader.metadata['/Title']
                    if title:
                        return title.strip()
                # Try first page text
                if len(reader.pages) > 0:
                    page = reader.pages[0]
                    text = page.extract_text()
                    if text:
                        # Get first line or first sentence
                        lines = text.split('\n')
                        for line in lines[:3]:
                            line = line.strip()
                            if len(line) > 10 and len(line) < 200:
                                return line
        except ImportError:
            pass
        
        # Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                if pdf.metadata and pdf.metadata.get('Title'):
                    return pdf.metadata['Title'].strip()
                # Try first page
                if len(pdf.pages) > 0:
                    page = pdf.pages[0]
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')
                        for line in lines[:3]:
                            line = line.strip()
                            if len(line) > 10 and len(line) < 200:
                                return line
        except ImportError:
            pass
    except Exception as e:
        print(f"Error reading PDF {pdf_path.name}: {e}")
    
    return None

def extract_docx_title(docx_path):
    """Try to extract title from DOCX"""
    try:
        from docx import Document
        doc = Document(docx_path)
        
        # Try to get title from first paragraph
        if len(doc.paragraphs) > 0:
            first_para = doc.paragraphs[0].text.strip()
            if len(first_para) > 10 and len(first_para) < 200:
                return first_para
        
        # Try core properties
        if doc.core_properties.title:
            return doc.core_properties.title.strip()
    except ImportError:
        pass
    except Exception as e:
        print(f"Error reading DOCX {docx_path.name}: {e}")
    
    return None

def main():
    target_dir = Path("/Users/macbookpro/Git Actions/flutter-opad/web/assets/uploads")
    data_file = Path("/Users/macbookpro/Git Actions/flutter-opad/lib/data/uploaded_files_data.dart")
    
    print("=" * 70)
    print("GENERATING TITLES FOR FILES")
    print("=" * 70)
    
    if not target_dir.exists():
        print("⚠️  Uploads directory not found")
        return False
    
    # Read current data file
    with open(data_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find all file entries
    pattern = r"name:\s*'([^']+)',\s*path:\s*'([^']+)'"
    matches = re.findall(pattern, content)
    
    print(f"\nFound {len(matches)} files to process")
    
    titles = {}
    improved_count = 0
    
    for filename, filepath in matches:
        file_path = target_dir / filepath
        
        if not file_path.exists():
            # Use improved filename
            title = improve_filename(filename)
            titles[filepath] = title
            continue
        
        # Try to extract title from file content
        title = None
        
        if file_path.suffix.lower() == '.pdf':
            title = extract_pdf_title(file_path)
        elif file_path.suffix.lower() in ['.docx']:
            title = extract_docx_title(file_path)
        elif file_path.suffix.lower() == '.doc':
            # Old DOC format - just improve filename
            title = None
        
        # Fallback to improved filename
        if not title:
            title = improve_filename(filename)
        else:
            improved_count += 1
        
        titles[filepath] = title
        print(f"  {filename[:50]:<50} -> {title[:60]}")
    
    # Update data file with titles
    # We'll add a displayTitle field or update the name field
    # For now, let's update the displayName logic in the model
    
    print(f"\n{'=' * 70}")
    print(f"Generated {len(titles)} titles")
    print(f"Extracted from content: {improved_count}")
    print(f"Improved from filename: {len(titles) - improved_count}")
    print("=" * 70)
    
    # Save titles to a separate file for reference
    titles_file = Path("/Users/macbookpro/Git Actions/flutter-opad/lib/data/file_titles.dart")
    
    dart_code = '''// File Titles
// Generated titles for better display in FilesScreen

class FileTitles {
  static final Map<String, String> titles = {
'''
    
    for filepath, title in titles.items():
        filepath_escaped = filepath.replace("'", "\\'")
        title_escaped = title.replace("'", "\\'")
        dart_code += f"    '{filepath_escaped}': '{title_escaped}',\n"
    
    dart_code += '''  };

  static String? getTitle(String path) {
    return titles[path];
  }
}
'''
    
    with open(titles_file, 'w', encoding='utf-8') as f:
        f.write(dart_code)
    
    print(f"\n✓ Saved titles to {titles_file}")
    print("\nNext: Update UploadedFile model to use these titles")
    
    return True

if __name__ == "__main__":
    main()

